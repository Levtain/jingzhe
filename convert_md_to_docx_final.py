#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def add_formatted_text(paragraph, text):
    """解析markdown格式并添加到段落"""
    # 解析粗体 **text**
    parts = re.split(r'\*\*(.+?)\*\*', text)

    for i, part in enumerate(parts):
        if i % 2 == 1:  # 奇数索引是粗体内容
            paragraph.add_run(part).bold = True
        else:
            # 普通文本，可能包含其他格式
            paragraph.add_run(part)

def convert_md_to_docx(md_path, output_path):
    """将markdown文件转换为docx格式"""

    # 读取markdown文件
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # 创建Word文档
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)

    # 设置标题样式
    for level in [1, 2, 3, 4]:
        heading = doc.styles[f'Heading {level}']
        heading.font.name = '微软雅黑'
        heading.font.size = Pt(18 - level * 2)  # 16, 14, 12, 10
        heading.font.bold = True
        heading.font.color.rgb = RGBColor(0, 0, 0)

    # 分割内容
    lines = content.split('\n')
    i = 0

    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # 处理代码块
        if line.startswith('```'):
            if in_code_block:
                # 代码块结束
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(50, 50, 50)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 空行跳过
        if not line.strip():
            i += 1
            continue

        # 处理标题
        if line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 处理列表
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, line.strip()[2:])
        elif re.match(r'^\d+\.', line.strip()):
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, line.strip())

        # 处理表格
        elif line.startswith('|'):
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1

            # 解析表格
            if len(table_lines) > 1:
                rows = []
                for table_line in table_lines:
                    if '|---' in table_line:
                        continue  # 跳过分隔行
                    cols = [c.strip() for c in table_line.split('|')]
                    cols = [c for c in cols if c]
                    if cols:
                        rows.append(cols)

                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'

                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx].cells):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_data

                                # 设置字体
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(10)

                                # 表头背景色
                                if row_idx == 0:
                                    from docx.oxml import OxmlElement
                                    shading_elm = OxmlElement('w:shd')
                                    shading_elm.set(qn('w:fill'), 'D5E8F0')
                                    cell._element.get_or_add_tcPr().append(shading_elm)
            continue

        # 处理普通段落
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line.strip())

        i += 1

    # 保存文档
    doc.save(output_path)
    print(f'Word文档已生成: {output_path}')
    return True

if __name__ == '__main__':
    md_file = r'd:\Claude\development\teaching\gamejam-planning-lecture-outline.md'
    output_file = r'C:\Users\Administrator\Desktop\GameJam策划讲座大纲.docx'

    convert_md_to_docx(md_file, output_file)
