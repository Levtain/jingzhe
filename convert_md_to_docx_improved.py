#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_formatted_text(paragraph, text):
    """解析markdown格式并添加到段落"""
    # 解析粗体 **text**
    parts = re.split(r'\*\*(.+?)\*\*', text)

    for i, part in enumerate(parts):
        if i % 2 == 1:  # 奇数索引是粗体内容
            paragraph.add_run(part).bold = True
        else:
            # 普通文本，可能包含其他格式
            paragraph.add_run(part)

def parse_list_item(line):
    """解析列表项,返回(层级, 文本, 是否编号)"""
    stripped = line.strip()

    # 无序列表
    if stripped.startswith('- ') or stripped.startswith('* '):
        return (0, stripped[2:], False)
    if re.match(r'^[\-\*]\s+', stripped):
        return (0, re.sub(r'^[\-\*]\s+', '', stripped), False)

    # 有序列表
    match = re.match(r'^(\d+)\.\s+(.+)', stripped)
    if match:
        return (0, match.group(2), True)

    # 缩进列表
    indent_match = re.match(r'^(\s+)[\-\*]\s+(.+)', stripped)
    if indent_match:
        indent_level = len(indent_match.group(1)) // 2
        return (indent_level, indent_match.group(2), False)

    indent_match = re.match(r'^(\s+)(\d+)\.\s+(.+)', stripped)
    if indent_match:
        indent_level = len(indent_match.group(1)) // 2
        return (indent_level, indent_match.group(3), True)

    return None

def is_code_block_line(line):
    """判断是否是代码块中的内容"""
    # 代码块通常包含特殊字符或格式
    code_indicators = [
        '│', '┌', '└', '─', '←', '→',
        '[ ]', '[x]', '✅', '❌',
        '→', 'P0:', 'P1:', 'P2:',
        'Day 1:', 'Day 2:',
    ]
    return any(indicator in line for indicator in code_indicators)

def convert_md_to_docx(md_path, output_path):
    """将markdown文件转换为docx格式"""

    # 读取markdown文件
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # 创建Word文档
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)

    # 设置标题样式
    for level in [1, 2, 3, 4]:
        heading = doc.styles[f'Heading {level}']
        heading.font.name = '微软雅黑'
        heading.font.size = Pt(20 - level * 2)
        heading.font.bold = True
        heading.font.color.rgb = RGBColor(0, 0, 0)

    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ''

    while i < len(lines):
        line = lines[i].rstrip('\n\r')

        # 处理代码块
        if line.startswith('```'):
            if in_code_block:
                # 代码块结束 - 检查是否是表格
                if code_lines and any('|' in line for line in code_lines):
                    # 作为文本段落处理
                    for code_line in code_lines:
                        if code_line.strip():
                            p = doc.add_paragraph()
                            p.paragraph_format.left_indent = Inches(0.5)
                            add_formatted_text(p, code_line)
                else:
                    # 作为代码块处理
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(50, 50, 50)

                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 空行跳过
        if not line.strip():
            i += 1
            continue

        # 处理分隔线
        if line.strip() == '---':
            doc.add_page_break()
            i += 1
            continue

        # 处理标题
        if line.startswith('#### '):
            title = re.sub(r'^####\s+\*\*(.+?)\*\*', r'\1', line)  # 移除标题中的**
            title = re.sub(r'^####\s+', '', title)
            p = doc.add_heading(title, level=4)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('### '):
            title = re.sub(r'^###\s+\*\*(.+?)\*\*', r'\1', line)  # 移除标题中的**
            title = re.sub(r'^###\s+', '', title)
            p = doc.add_heading(title, level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            title = re.sub(r'^##\s+\*\*(.+?)\*\*', r'\1', line)  # 移除标题中的**
            title = re.sub(r'^##\s+', '', title)
            p = doc.add_heading(title, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('# '):
            title = re.sub(r'^#\s+\*\*(.+?)\*\*', r'\1', line)  # 移除标题中的**
            title = re.sub(r'^#\s+', '', title)
            p = doc.add_heading(title, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 处理列表
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, line.strip()[2:])
        elif re.match(r'^\s*[\-\*]\s+', line):
            # 缩进列表
            indent_level = len(re.match(r'^(\s*)', line).group(1)) // 2
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            text = re.sub(r'^\s*[\-\*]\s+', '', line)
            add_formatted_text(p, text)
        elif re.match(r'^\s*\d+\.\s+', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            add_formatted_text(p, text)

        # 处理表格
        elif line.startswith('|'):
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i].rstrip('\n\r'))
                i += 1

            # 解析表格
            if len(table_lines) > 1:
                rows = []
                for table_line in table_lines:
                    if '|---' in table_line or '|===' in table_line:
                        continue  # 跳过分隔行
                    cols = [c.strip() for c in table_line.split('|')]
                    cols = [c for c in cols if c]
                    if cols:
                        rows.append(cols)

                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'

                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx].cells):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_data

                                # 设置字体
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(10)

                                # 表头背景色
                                if row_idx == 0:
                                    shading_elm = OxmlElement('w:shd')
                                    shading_elm.set(qn('w:fill'), 'D5E8F0')
                                    cell._element.get_or_add_tcPr().append(shading_elm)
            continue

        # 处理普通段落
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line.strip())

        i += 1

    # 保存文档
    doc.save(output_path)
    print(f'Word文档已生成: {output_path}')
    return True

if __name__ == '__main__':
    md_file = r'd:\Claude\development\teaching\gamejam-planning-lecture-outline.md'
    output_file = r'C:\Users\Administrator\Desktop\GameJam策划讲座大纲.docx'

    convert_md_to_docx(md_file, output_file)
